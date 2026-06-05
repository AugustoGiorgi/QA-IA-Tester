from typing import List, Dict
from docx import Document


def docx_to_text(path: str) -> str:
    """
    Extrae un texto lineal del DOCX (párrafos + tablas a modo 'fila | fila').
    Se usa para heurísticas de otras secciones, headings, etc.
    """
    doc = Document(path)
    texts = []

    # Procesar párrafos
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            texts.append(p.text.strip())

    # Procesar tablas (como líneas legibles)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_txt = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_txt:
                texts.append(" | ".join(row_txt))

    return "\n".join(texts)


# === Tablas estructuradas (para evaluaciones específicas) ===
def extract_tables(docx_path: str) -> List[Dict]:
    """
    Devuelve una lista de tablas con forma:
    [
      {"headers": ["Col1", "Col2", ...],
       "rows": [["v11","v12",...], ["v21","v22",...], ...]
      },
      ...
    ]
    La primera fila se toma como encabezados si existe.
    """
    doc = Document(docx_path)
    out: List[Dict] = []

    for t in doc.tables:
        headers = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        rows = []
        for r in t.rows[1:]:
            rows.append([c.text.strip() for c in r.cells])
        out.append({"headers": headers, "rows": rows})

    return out
