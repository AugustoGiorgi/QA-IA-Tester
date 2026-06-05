# backend/services/generator.py
import json
import re
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

from services.ai import build_messages, complete

GEN_PROMPT = (
    "Actuá como analista funcional senior. Generá un DOCUMENTO FUNCIONAL de alta calidad, "
    "optimizado para obtener un puntaje >= 90% según criterios de QA: claridad, completitud, "
    "flujos principal/alternos, casos borde, reglas de negocio, datos de prueba ejemplo, "
    "criterios de aceptación medibles, dependencias, supuestos, riesgos, alcance y fuera de alcance, "
    "requisitos NO funcionales (seguridad, performance, observabilidad). "
    "Devolvé SOLO un JSON válido con esta estructura (ejemplo de keys, sin texto extra):\n"
    "{\n"
    '  "Titulo": "string",\n'
    '  "ResumenEjecutivo": "string",\n'
    '  "Objetivos": ["string", "..."],\n'
    '  "Alcance": ["in-scope 1", "in-scope 2"],\n'
    '  "FueraDeAlcance": ["out 1", "out 2"],\n'
    '  "Actores": ["Usuario", "Admin", "Integración X"],\n'
    '  "HistoriasDeUsuario": [\n'
    '    {"ID":"HU-1","Como":"<actor>","Quiero":"<objetivo>","Para":"<beneficio>", "CriteriosDeAceptacion":["..."]}\n'
    "  ],\n"
    '  "ReglasDeNegocio": ["regla 1", "regla 2"],\n'
    '  "Flujos": {"Principal": ["paso 1", "paso 2"], "Alternativos": [["alt-1 paso 1","alt-1 paso 2"]]},\n'
    '  "CasosBorde": ["borde 1", "borde 2"],\n'
    '  "DatosDePrueba": ["ejemplo1=...", "ejemplo2=..."],\n'
    '  "CriteriosDeAceptacionGlobales": ["medible 1", "medible 2"],\n'
    '  "RequisitosNoFuncionales": ["seguridad ...", "performance ...", "observabilidad ..."],\n'
    '  "Dependencias": ["sistema A", "API B"],\n'
    '  "Integraciones": ["servicio X", "servicio Y"],\n'
    '  "Supuestos": ["..."],\n'
    '  "Riesgos": ["..."]\n'
    "}\n"
    "Buenas prácticas: evitar ambigüedades, usar lenguaje medible (NFR con umbrales)."
)

def _safe_json_loads(text: str) -> Dict[str, Any]:
    # intenta extraer el primer bloque JSON válido
    m = re.search(r"\{.*\}", text, flags=re.S)
    if not m:
        return {}
    candidate = m.group(0)
    try:
        return json.loads(candidate)
    except Exception:
        return {}

def generate_functional_json(context: Dict[str, Any], feedback_snippets: Optional[List[str]] = None) -> Dict[str, Any]:
    ctx_lines = []
    for k, v in context.items():
        if v is None:
            continue
        if isinstance(v, list):
            ctx_lines.append(f"{k}: " + ", ".join([str(x) for x in v]))
        else:
            ctx_lines.append(f"{k}: {v}")
    ctx = "\n".join(ctx_lines)

    messages = build_messages(GEN_PROMPT, f"Contexto de negocio/alcance:\n{ctx}", feedback_snippets)
    resp = complete(messages, temperature=0.1)
    data = _safe_json_loads(resp)
    return data or {
        "Titulo": context.get("titulo", "Documento Funcional"),
        "ResumenEjecutivo": context.get("contexto", "") or "Resumen no provisto.",
        "Objetivos": [context.get("objetivos", "")] if context.get("objetivos") else [],
        "Alcance": [context.get("alcance", "")] if context.get("alcance") else [],
        "Actores": context.get("actores", []) or [],
    }

def _add_heading(doc: Document, text: str, level: int):
    run = doc.add_heading(level=level).add_run(text)
    run.font.size = Pt(14 if level >= 2 else 16)

def _add_list(doc: Document, items: List[Any]):
    for it in items or []:
        p = doc.add_paragraph(str(it))
        p.style = "List Bullet"

def _add_table_hu(doc: Document, historias: List[Dict[str, Any]]):
    if not historias:
        return
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    headers = ["ID", "Como", "Quiero", "Para", "Criterios de aceptación"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for hu in historias:
        row = table.add_row().cells
        row[0].text = str(hu.get("ID", "HU"))
        row[1].text = str(hu.get("Como", ""))
        row[2].text = str(hu.get("Quiero", ""))
        row[3].text = str(hu.get("Para", ""))
        ca = hu.get("CriteriosDeAceptacion", []) or []
        row[4].text = "\n".join([f"- {x}" for x in ca])

def functional_json_to_docx(data: Dict[str, Any], out_path: str):
    doc = Document()
    # Título
    doc.add_heading(data.get("Titulo", "Documento Funcional"), 0)

    # Resumen
    if data.get("ResumenEjecutivo"):
        _add_heading(doc, "Resumen ejecutivo", 1)
        doc.add_paragraph(str(data["ResumenEjecutivo"]))

    # Objetivos & Alcance
    if data.get("Objetivos"):
        _add_heading(doc, "Objetivos", 1)
        _add_list(doc, data["Objetivos"])
    if data.get("Alcance"):
        _add_heading(doc, "Alcance (In-scope)", 1)
        _add_list(doc, data["Alcance"])
    if data.get("FueraDeAlcance"):
        _add_heading(doc, "Fuera de alcance (Out-of-scope)", 1)
        _add_list(doc, data["FueraDeAlcance"])

    # Actores
    if data.get("Actores"):
        _add_heading(doc, "Actores", 1)
        _add_list(doc, data["Actores"])

    # Historias
    if data.get("HistoriasDeUsuario"):
        _add_heading(doc, "Historias de usuario", 1)
        _add_table_hu(doc, data["HistoriasDeUsuario"])

    # Reglas de negocio
    if data.get("ReglasDeNegocio"):
        _add_heading(doc, "Reglas de negocio", 1)
        _add_list(doc, data["ReglasDeNegocio"])

    # Flujos
    flujos = data.get("Flujos") or {}
    if flujos:
        _add_heading(doc, "Flujos", 1)
        principal = flujos.get("Principal") or []
        if principal:
            _add_heading(doc, "Flujo principal", 2)
            for idx, paso in enumerate(principal, 1):
                doc.add_paragraph(f"{idx}. {paso}")
        alternos = flujos.get("Alternativos") or []
        if alternos:
            _add_heading(doc, "Flujos alternativos", 2)
            for alt in alternos:
                doc.add_paragraph().add_run("\n").add_break(WD_BREAK.LINE)
                for idx, paso in enumerate(alt, 1):
                    doc.add_paragraph(f"{idx}. {paso}")

    # Casos borde
    if data.get("CasosBorde"):
        _add_heading(doc, "Casos borde", 1)
        _add_list(doc, data["CasosBorde"])

    # Datos de prueba
    if data.get("DatosDePrueba"):
        _add_heading(doc, "Datos de prueba (ejemplos)", 1)
        _add_list(doc, data["DatosDePrueba"])

    # Criterios de aceptación globales
    if data.get("CriteriosDeAceptacionGlobales"):
        _add_heading(doc, "Criterios de aceptación (globales)", 1)
        _add_list(doc, data["CriteriosDeAceptacionGlobales"])

    # NFR
    if data.get("RequisitosNoFuncionales"):
        _add_heading(doc, "Requisitos NO funcionales", 1)
        _add_list(doc, data["RequisitosNoFuncionales"])

    # Dependencias & Integraciones
    if data.get("Dependencias"):
        _add_heading(doc, "Dependencias", 1)
        _add_list(doc, data["Dependencias"])
    if data.get("Integraciones"):
        _add_heading(doc, "Integraciones", 1)
        _add_list(doc, data["Integraciones"])

    # Supuestos & Riesgos
    if data.get("Supuestos"):
        _add_heading(doc, "Supuestos", 1)
        _add_list(doc, data["Supuestos"])
    if data.get("Riesgos"):
        _add_heading(doc, "Riesgos", 1)
        _add_list(doc, data["Riesgos"])

    doc.save(out_path)
